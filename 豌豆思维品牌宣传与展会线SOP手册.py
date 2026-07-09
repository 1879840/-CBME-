#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""豌豆思维品牌宣传与展会线SOP手册生成器"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

print("开始生成豌豆思维SOP手册...")

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Calibri'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)

# ========== 封面 ==========
title = doc.add_heading('豌豆思维', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(76, 175, 80)

subtitle = doc.add_heading('品牌宣传与展会线SOP手册', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(22)

doc.add_paragraph('\n\n')
p1 = doc.add_paragraph('版本：V1.0')
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2 = doc.add_paragraph('日期：2026年6月')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3 = doc.add_paragraph('适用范围：品牌宣传岗、专家线、展会线、市场部')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ========== 第一部分 ==========
doc.add_heading('第一部分：素质教育行业品牌宣传岗位SOP', level=1)

doc.add_heading('一、行业特性（决定SOP的底层逻辑）', level=2)
doc.add_paragraph('素质教育与思维教育行业具有以下显著特征：')
doc.add_paragraph('• 决策者≠使用者：付费方是家长（多为85后/90后宝妈），上课方是孩子，宣传需要"双轨表达"')
doc.add_paragraph('• 强信任驱动：客单价高（年课2000–20000元）、决策周期长（平均7–30天），品牌内容承担"种草+培育"双重任务')
